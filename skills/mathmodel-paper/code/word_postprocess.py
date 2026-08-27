#!/usr/bin/env python3
"""word_postprocess.py — pandoc 转换后的 Word 版式微调（python-docx）。

用法：
    python3 code/word_postprocess.py paper/paper.docx
    python3 code/word_postprocess.py            # 默认 paper/paper.docx

**角色定位（重要）**：Word 文档由 pandoc 从 paper.tex 转换生成，公式已自动转为 OMML
原生数学格式。本脚本**仅用于 pandoc 转换后的版式微调**（页眉留空、页码页脚、中文字体、
三线表核对），**禁止用它从头生成全文或重建内容**（不得 add_paragraph / add_table /
add_page_break 新增任何段落或表格），**禁止手工插入公式**，否则 Word 中将丢失全部数学公式。

来源：主技能 skills/math-modeling-helper/SKILL.md
      「Word 文档后处理模块（python-docx，用于 pandoc 转换后的版式微调）」小节。
"""
import argparse

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cjk_font(run, font_name='宋体', size=Pt(12), bold=False):
    """统一设置中西文字体与字号（先设西文名确保 rPr/rFonts 存在，再设中文名）"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    run.bold = bold


def set_three_line_table(table):
    """设置三线表样式：顶线粗、底线粗、表头下线细，无竖线"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # 删除默认边框
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'bottom', 'left', 'right', 'insideV', 'insideH'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'none')
        elem.set(qn('w:sz'), '0')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), '000000')
        borders.append(elem)
    # 顶线（粗 1.5pt）
    top = borders.find(qn('w:top'))
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    # 底线（粗 1.5pt）
    bottom = borders.find(qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    tblPr.append(borders)
    # 表头行底部加细线（0.5pt）
    if len(table.rows) > 1:
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            cell_borders = OxmlElement('w:tcBorders')
            btm = OxmlElement('w:bottom')
            btm.set(qn('w:val'), 'single')
            btm.set(qn('w:sz'), '4')
            btm.set(qn('w:space'), '0')
            btm.set(qn('w:color'), '000000')
            cell_borders.append(btm)
            tcPr.append(cell_borders)


def set_cell_center(cell, text=None, font_name='宋体', font_size=Pt(12)):
    """设置单元格文字上下左右居中（text=None 时仅调整既有内容样式，不重建内容）"""
    if text is not None:
        cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中
    if text is None:
        for r in p.runs:
            set_cjk_font(r, font_name, font_size)
    else:
        run = p.add_run(text)
        set_cjk_font(run, font_name, font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 垂直居中


def restyle_word_paper(doc_path='../paper/paper.docx'):
    """pandoc 转换后的版式微调：仅调整既有内容的样式（页边距/页眉页脚/字体/三线表），
    禁止 add_paragraph/add_table 新增或重建任何内容，否则会在 pandoc 文档后追加重复全文"""
    doc = Document(doc_path)  # 读取 pandoc 转换结果

    # ---- 页面设置：A4，四边2.5cm；页眉留空；页脚居中页码 ----
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        # 页眉必须留空（匿名规范）
        for p in section.header.paragraphs:
            p.text = ''
        # 页脚：居中页码（插入 PAGE 域，Word 打开时自动更新）
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = ''
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'PAGE'
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        for el in (fld_begin, instr, fld_end):
            run._element.append(el)

    # ---- 遍历既有段落，按样式/内容特征微调（不新增任何段落） ----
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style is not None else ''
        if style_name.startswith('Title'):            # 论文题目：三号黑体居中
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_cjk_font(r, '黑体', Pt(16), bold=True)
        elif style_name.startswith('Heading 1'):      # 一级标题：四号黑体居中
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_cjk_font(r, '黑体', Pt(14), bold=True)
        elif text.startswith('摘要'):                  # 摘要标签段：四号黑体
            for r in para.runs:
                set_cjk_font(r, '黑体', Pt(14), bold=True)
        else:                                          # 正文：小四宋体，1倍行距，首行缩进2字符
            para.paragraph_format.line_spacing = 1
            para.paragraph_format.first_line_indent = Cm(0.85)
            for r in para.runs:
                set_cjk_font(r, '宋体', Pt(12))

    # ---- 既有表格统一三线表样式 + 单元格文字居中（不新建表格） ----
    for table in doc.tables:
        set_three_line_table(table)
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 保存（覆盖写回 pandoc 结果）
    doc.save(doc_path)
    print(f'Word 版式微调完成：{doc_path}')


def main():
    parser = argparse.ArgumentParser(
        description='pandoc 转换后的 Word 版式微调（仅改样式，禁止重建内容/手工插入公式）'
    )
    parser.add_argument(
        'docx', nargs='?', default='paper/paper.docx',
        help='待微调的 .docx 路径（pandoc 转换结果，原地覆盖写回），默认 paper/paper.docx'
    )
    args = parser.parse_args()
    restyle_word_paper(args.docx)


if __name__ == '__main__':
    main()
